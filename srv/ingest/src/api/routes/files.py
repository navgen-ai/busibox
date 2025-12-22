"""
File metadata, download, deletion, chunk browsing, reprocessing, and export endpoints.

Handles:
- GET /files/{fileId}: Retrieve file metadata
- GET /files/{fileId}/download: Download original file from MinIO
- GET /files/{fileId}/chunks: Retrieve text chunks for a file
- POST /files/{fileId}/search: Search within a single document
- DELETE /files/{fileId}: Delete file and all associated data
- POST /files/{fileId}/reprocess: Reprocess document (delete chunks/vectors, re-run ingestion)
- GET /files/{fileId}/export: Export document in various formats (markdown, html, text, docx, pdf)
"""

import io
import json
import uuid
from typing import Optional, List

import redis.asyncio as redis_async
import structlog
from fastapi import APIRouter, Depends, Request, Query, status
from fastapi.responses import JSONResponse, StreamingResponse, Response
from pydantic import BaseModel, Field

from api.middleware.jwt_auth import ScopeChecker
from api.services.minio_service import MinIOService
from api.services.postgres import PostgresService
from api.services.encryption_client import EncryptionClient
from services.milvus_service import MilvusService
from services.processing_history_service import ProcessingHistoryService
from shared.config import Config

logger = structlog.get_logger()

router = APIRouter()

# Scope dependencies
require_ingest_read = ScopeChecker("ingest.read")
require_ingest_write = ScopeChecker("ingest.write")
require_ingest_delete = ScopeChecker("ingest.delete")


def validate_uuid(file_id_str: str) -> tuple[Optional[uuid.UUID], Optional[JSONResponse]]:
    """
    Validate a string as a UUID.
    
    Returns:
        tuple: (uuid.UUID, None) if valid, or (None, JSONResponse) with 400 error if invalid
    """
    try:
        return uuid.UUID(file_id_str), None
    except ValueError:
        return None, JSONResponse(
            status_code=status.HTTP_400_BAD_REQUEST,
            content={"error": "Invalid file ID format", "details": "File ID must be a valid UUID"}
        )


async def check_file_access(
    conn,
    file_id: uuid.UUID,
    user_id: str,
    request: Request
) -> tuple[bool, Optional[dict], Optional[str]]:
    """
    Check if user has access to a file based on visibility and roles.
    
    Args:
        conn: Database connection
        file_id: File UUID
        user_id: User ID
        request: Request object with role_ids_read in state
    
    Returns:
        (has_access, file_row, error_message)
    """
    # Get file with visibility and role information
    # Only query essential columns that exist in all contexts
    file_row = await conn.fetchrow("""
        SELECT 
            f.file_id,
            f.user_id,
            f.owner_id,
            f.visibility,
            COALESCE(
                (SELECT json_agg(role_id) 
                 FROM document_roles 
                 WHERE file_id = f.file_id),
                '[]'::json
            ) as role_ids
        FROM ingestion_files f
        WHERE f.file_id = $1
    """, file_id)
    
    if not file_row:
        return False, None, "File not found"
    
    # Convert to dict for easier access
    file_dict = dict(file_row)
    visibility = file_dict.get("visibility", "personal")
    
    # Parse role_ids - it comes back as JSON from the query
    import json
    role_ids_raw = file_dict.get("role_ids", [])
    if isinstance(role_ids_raw, str):
        file_role_ids = json.loads(role_ids_raw)
    elif isinstance(role_ids_raw, list):
        file_role_ids = role_ids_raw
    else:
        file_role_ids = []
    
    owner_id = str(file_dict.get("owner_id", ""))
    
    # Personal files: only owner can access
    if visibility == "personal":
        if user_id == owner_id:
            return True, file_dict, None
        else:
            return False, file_dict, "Unauthorized: personal file owned by another user"
    
    # Shared files: check role intersection
    if visibility == "shared":
        # Get user's role IDs from request state
        # Note: JWT middleware sets "role_ids" not "role_ids_read"
        user_role_ids = getattr(request.state, "role_ids", [])
        
        # Check if user has at least one matching role
        if not file_role_ids:
            # Shared file with no roles - only owner can access
            if user_id == owner_id:
                return True, file_dict, None
            else:
                return False, file_dict, "Unauthorized: shared file with no roles"
        
        # Check for role intersection
        file_role_set = set(str(r) for r in file_role_ids)
        user_role_set = set(str(r) for r in user_role_ids)
        
        if file_role_set & user_role_set:
            # User has at least one matching role
            return True, file_dict, None
        else:
            return False, file_dict, f"Unauthorized: file requires roles {file_role_set}, user has {user_role_set}"
    
    # Unknown visibility - deny access
    return False, file_dict, f"Unknown visibility: {visibility}"


@router.get("/{fileId}", dependencies=[Depends(require_ingest_read)])
async def get_file_metadata(fileId: str, request: Request):
    """
    Get file metadata and current status.
    
    Returns:
        File metadata including status, processing metrics, extracted metadata
    """
    # Validate UUID format
    file_uuid, error_response = validate_uuid(fileId)
    if error_response:
        return error_response
    
    user_id = request.state.user_id
    
    from api.main import pg_service as postgres_service  # Use shared PostgresService instance
    # Connection is already established in startup
    
    try:
        async with postgres_service.acquire(request) as conn:
            # Check file access based on visibility and roles
            has_access, file_data, error_msg = await check_file_access(
                conn, file_uuid, user_id, request
            )
            
            if not has_access:
                status_code = status.HTTP_404_NOT_FOUND if error_msg == "File not found" else status.HTTP_403_FORBIDDEN
                return JSONResponse(
                    status_code=status_code,
                    content={"error": error_msg}
                )
            
            # Get full file record with all metadata
            file_row = await conn.fetchrow("""
                SELECT 
                    f.file_id,
                    f.user_id,
                    f.owner_id,
                    f.visibility,
                    f.filename,
                    f.original_filename,
                    f.mime_type,
                    f.size_bytes,
                    f.storage_path,
                    f.content_hash,
                    f.document_type,
                    f.primary_language,
                    f.detected_languages,
                    f.classification_confidence,
                    f.chunk_count,
                    f.vector_count,
                    f.processing_duration_seconds,
                    f.extracted_title,
                    f.extracted_author,
                    f.extracted_date,
                    f.extracted_keywords,
                    f.metadata,
                    f.permissions,
                    f.created_at,
                    f.updated_at
                FROM ingestion_files f
                WHERE f.file_id = $1
            """, uuid.UUID(fileId))
            
            # Get status
            status_row = await conn.fetchrow("""
                SELECT 
                    stage, progress, chunks_processed, total_chunks,
                    pages_processed, total_pages, error_message,
                    started_at, completed_at, updated_at
                FROM ingestion_status
                WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            # Get processing strategies attempted
            strategy_rows = await conn.fetch("""
                SELECT 
                    processing_strategy,
                    success,
                    text_length,
                    chunk_count,
                    embedding_count,
                    visual_embedding_count,
                    processing_time_seconds,
                    error_message,
                    metadata,
                    created_at
                FROM processing_strategy_results
                WHERE file_id = $1
                ORDER BY created_at ASC
            """, uuid.UUID(fileId))
            
            # Get page_count and word_count from processing_history metadata (fallback)
            history_metadata_row = await conn.fetchrow("""
                SELECT metadata
                FROM processing_history
                WHERE file_id = $1 
                  AND metadata IS NOT NULL 
                  AND (metadata->>'page_count' IS NOT NULL OR metadata->>'text_length' IS NOT NULL)
                ORDER BY created_at DESC
                LIMIT 1
            """, uuid.UUID(fileId))
            
            strategies = [
                {
                    "strategy": row["processing_strategy"],
                    "success": row["success"],
                    "textLength": row["text_length"],
                    "chunkCount": row["chunk_count"],
                    "embeddingCount": row["embedding_count"],
                    "visualEmbeddingCount": row["visual_embedding_count"],
                    "processingTimeSeconds": float(row["processing_time_seconds"]) if row["processing_time_seconds"] else None,
                    "errorMessage": row["error_message"],
                    "metadata": row["metadata"],
                    "attemptedAt": row["created_at"].isoformat() if row["created_at"] else None,
                }
                for row in strategy_rows
            ]
            
            # Merge metadata with fallbacks for page_count and word_count
            # Handle case where metadata might be a JSON string instead of dict
            def ensure_dict(val):
                if val is None:
                    return {}
                if isinstance(val, str):
                    try:
                        return json.loads(val)
                    except (json.JSONDecodeError, TypeError):
                        return {}
                return val if isinstance(val, dict) else {}
            
            metadata = ensure_dict(file_row["metadata"])
            history_meta = ensure_dict(history_metadata_row["metadata"] if history_metadata_row else None)
            
            # Fallback for page_count from multiple sources
            if "page_count" not in metadata or metadata["page_count"] is None:
                # Try processing history first (most reliable)
                if history_meta and history_meta.get("page_count"):
                    metadata["page_count"] = history_meta["page_count"]
                # Then try status table
                elif status_row and status_row["total_pages"]:
                    metadata["page_count"] = status_row["total_pages"]
                # Finally try strategy metadata
                elif strategies:
                    for strategy in strategies:
                        if strategy.get("success") and strategy.get("metadata"):
                            strat_meta = ensure_dict(strategy["metadata"])
                            if strat_meta.get("page_count"):
                                metadata["page_count"] = strat_meta["page_count"]
                                break
            
            # Fallback for word_count from multiple sources
            if "word_count" not in metadata or metadata["word_count"] is None:
                # Try processing history text_length first
                if history_meta and history_meta.get("text_length"):
                    # Estimate words from characters (average 5 chars per word)
                    metadata["word_count"] = history_meta["text_length"] // 5
                else:
                    # Try to calculate from successful strategy text_length
                    for strategy in strategies:
                        if strategy.get("success") and strategy.get("textLength"):
                            metadata["word_count"] = strategy["textLength"] // 5
                            break
                # If still no word_count, estimate from chunk_count
                if ("word_count" not in metadata or metadata["word_count"] is None) and file_row["chunk_count"]:
                    # Rough estimate: ~200 words per chunk on average
                    metadata["word_count"] = file_row["chunk_count"] * 200
            
            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "fileId": str(file_row["file_id"]),
                    "filename": file_row["filename"],
                    "originalFilename": file_row["original_filename"],
                    "mimeType": file_row["mime_type"],
                    "sizeBytes": file_row["size_bytes"],
                    "contentHash": file_row["content_hash"],
                    "documentType": file_row["document_type"],
                    "primaryLanguage": file_row["primary_language"],
                    "detectedLanguages": file_row["detected_languages"],
                    "classificationConfidence": float(file_row["classification_confidence"]) if file_row["classification_confidence"] else None,
                    "chunkCount": file_row["chunk_count"],
                    "vectorCount": file_row["vector_count"],
                    "processingDurationSeconds": file_row["processing_duration_seconds"],
                    "extractedTitle": file_row["extracted_title"],
                    "extractedAuthor": file_row["extracted_author"],
                    "extractedDate": file_row["extracted_date"].isoformat() if file_row["extracted_date"] else None,
                    "extractedKeywords": file_row["extracted_keywords"],
                    "metadata": metadata,
                    "permissions": file_row["permissions"],
                    "processingStrategies": strategies,
                    "status": {
                        "stage": status_row["stage"] if status_row else None,
                        "progress": status_row["progress"] if status_row else None,
                        "chunksProcessed": status_row["chunks_processed"] if status_row else None,
                        "totalChunks": status_row["total_chunks"] if status_row else None,
                        "pagesProcessed": status_row["pages_processed"] if status_row else None,
                        "totalPages": status_row["total_pages"] if status_row else None,
                        "errorMessage": status_row["error_message"] if status_row else None,
                        "startedAt": status_row["started_at"].isoformat() if status_row and status_row["started_at"] else None,
                        "completedAt": status_row["completed_at"].isoformat() if status_row and status_row["completed_at"] else None,
                        "updatedAt": status_row["updated_at"].isoformat() if status_row and status_row["updated_at"] else None,
                    },
                    "createdAt": file_row["created_at"].isoformat(),
                    "updatedAt": file_row["updated_at"].isoformat(),
                }
            )
    
    except ValueError as e:
        # Handle UUID parsing errors (invalid file ID format)
        logger.warning(
            "Invalid file ID format",
            file_id=fileId,
            user_id=user_id,
            error=str(e),
        )
        return JSONResponse(
            status_code=status.HTTP_400_BAD_REQUEST,
            content={"error": "Invalid file ID format"}
        )
    
    except Exception as e:
        logger.error(
            "Failed to get file metadata",
            file_id=fileId,
            user_id=user_id,
            error=str(e),
            exc_info=True,
        )
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={"error": "Failed to retrieve file metadata", "details": str(e)}
        )
    
    finally:
        # Don't disconnect - pg_service is a singleton shared across requests
        pass


@router.get("/{fileId}/history", dependencies=[Depends(require_ingest_read)])
async def get_processing_history(fileId: str, request: Request):
    """
    Get detailed processing history for a file.
    
    Returns:
        Processing history with steps, timing, and any errors
    """
    user_id = request.state.user_id
    
    try:
        # Validate UUID format
        try:
            uuid.UUID(fileId)
        except ValueError:
            return JSONResponse(
                status_code=status.HTTP_400_BAD_REQUEST,
                content={"error": "Invalid file ID format"}
            )
        
        config = Config().to_dict()
        history_service = ProcessingHistoryService(config)
        history_service.connect()  # Must connect before use
        
        try:
            history = history_service.get_history(fileId)
            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={"history": history}
            )
        finally:
            history_service.disconnect()
            
    except Exception as e:
        logger.error(
            "Failed to get processing history",
            file_id=fileId,
            user_id=user_id,
            error=str(e),
            exc_info=True,
        )
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={"error": "Failed to retrieve processing history"}
        )
        pass


@router.get("/{fileId}/download", dependencies=[Depends(require_ingest_read)])
async def download_file(fileId: str, request: Request):
    """
    Download original file from MinIO storage.
    
    Content is automatically decrypted if envelope encryption was used.
    
    Returns:
        StreamingResponse with file content
    """
    # Validate UUID format
    file_uuid, error_response = validate_uuid(fileId)
    if error_response:
        return error_response
    
    user_id = request.state.user_id
    
    config = Config().to_dict()
    postgres_service = PostgresService(config, request)
    minio_service = MinIOService(config)
    encryption_client = EncryptionClient(config)
    
    await postgres_service.connect()
    
    try:
        async with postgres_service.acquire(request) as conn:
            # Check file access based on visibility and roles
            has_access, file_data, error_msg = await check_file_access(
                conn, file_uuid, user_id, request
            )
            
            if not has_access:
                status_code = status.HTTP_404_NOT_FOUND if error_msg == "File not found" else status.HTTP_403_FORBIDDEN
                return JSONResponse(
                    status_code=status_code,
                    content={"error": error_msg}
                )
            
            # Get file from MinIO using storage_path from file_data
            # Note: storage_path not in check_file_access, need to fetch it
            file_row = await conn.fetchrow("""
                SELECT storage_path, original_filename, mime_type, visibility
                FROM ingestion_files
                WHERE file_id = $1
            """, file_uuid)
            
            storage_path = file_row["storage_path"]
            original_filename = file_row["original_filename"]
            mime_type = file_row["mime_type"]
            visibility = file_row["visibility"]
            
            # Get file role IDs for decryption
            role_ids = getattr(request.state, "role_ids_read", [])
            
            # Get file from MinIO
            try:
                import asyncio
                loop = asyncio.get_event_loop()
                
                # Get file object from MinIO
                minio_response = await loop.run_in_executor(
                    None,
                    lambda: minio_service.client.get_object(
                        minio_service.bucket,
                        storage_path
                    )
                )
                
                # Read file content
                content = await loop.run_in_executor(None, minio_response.read)
                
                # Decrypt content if it appears to be encrypted
                if encryption_client.enabled and encryption_client.is_encrypted(content):
                    try:
                        content = await encryption_client.decrypt_for_download(
                            file_id=fileId,
                            encrypted_content=content,
                            role_ids=role_ids,
                            user_id=user_id if visibility == "personal" else None,
                        )
                        logger.info(
                            "File content decrypted for download",
                            file_id=fileId,
                            user_id=user_id,
                            decrypted_size=len(content),
                        )
                    except PermissionError as e:
                        logger.warning(
                            "Decryption access denied",
                            file_id=fileId,
                            user_id=user_id,
                            error=str(e),
                        )
                        return JSONResponse(
                            status_code=status.HTTP_403_FORBIDDEN,
                            content={"error": "Access denied to decrypt file"}
                        )
                
                logger.info(
                    "File downloaded",
                    file_id=fileId,
                    user_id=user_id,
                    filename=original_filename,
                    size_bytes=len(content),
                )
                
                # Return file as streaming response
                return StreamingResponse(
                    iter([content]),
                    media_type=mime_type,
                    headers={
                        'Content-Disposition': f'attachment; filename="{original_filename}"'
                    }
                )
                
            except Exception as e:
                logger.error(
                    "Failed to download file from MinIO",
                    file_id=fileId,
                    storage_path=storage_path,
                    error=str(e),
                    exc_info=True,
                )
                return JSONResponse(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    content={"error": "Failed to download file from storage", "details": str(e)}
                )
    
    except Exception as e:
        logger.error(
            "Failed to process download request",
            file_id=fileId,
            user_id=user_id,
            error=str(e),
            exc_info=True,
        )
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={"error": "Failed to process download request", "details": str(e)}
        )
    
    finally:
        # Don't disconnect - pg_service is a singleton shared across requests
        pass


@router.get("/{fileId}/presigned-url", dependencies=[Depends(require_ingest_read)])
async def get_presigned_url(fileId: str, request: Request, expiry: int = 3600):
    """
    Generate a presigned URL for direct file access from MinIO.
    
    Args:
        fileId: File identifier
        expiry: URL expiration time in seconds (default: 3600 = 1 hour)
    
    Returns:
        JSON with presigned URL
    """
    user_id = request.state.user_id
    
    config = Config().to_dict()
    postgres_service = PostgresService(config, request)
    minio_service = MinIOService(config)
    
    await postgres_service.connect()
    
    try:
        async with postgres_service.acquire(request) as conn:
            # Get file record
            file_row = await conn.fetchrow("""
                SELECT user_id, storage_path, original_filename, mime_type
                FROM ingestion_files
                WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            if not file_row:
                return JSONResponse(
                    status_code=status.HTTP_404_NOT_FOUND,
                    content={"error": "File not found"}
                )
            
            # Verify ownership
            if str(file_row["user_id"]) != user_id:
                return JSONResponse(
                    status_code=status.HTTP_403_FORBIDDEN,
                    content={"error": "Unauthorized access"}
                )
            
            storage_path = file_row["storage_path"]
            mime_type = file_row["mime_type"]
            filename = file_row["original_filename"]
            
            # Generate presigned URL
            try:
                import asyncio
                from datetime import timedelta
                
                loop = asyncio.get_event_loop()
                
                # Generate presigned GET URL
                presigned_url = await loop.run_in_executor(
                    None,
                    lambda: minio_service.client.presigned_get_object(
                        minio_service.bucket,
                        storage_path,
                        expires=timedelta(seconds=expiry)
                    )
                )
                
                logger.info(
                    "Presigned URL generated successfully",
                    file_id=fileId,
                    user_id=user_id,
                    storage_path=storage_path,
                    filename=filename,
                    mime_type=mime_type,
                    expiry_seconds=expiry,
                    presigned_url=presigned_url,
                    url_length=len(presigned_url),
                )
                
                return JSONResponse(
                    status_code=status.HTTP_200_OK,
                    content={
                        "url": presigned_url,
                        "expiresIn": expiry,
                    }
                )
                
            except Exception as e:
                logger.error(
                    "Failed to generate presigned URL",
                    file_id=fileId,
                    storage_path=storage_path,
                    error=str(e),
                    exc_info=True,
                )
                return JSONResponse(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    content={"error": "Failed to generate presigned URL", "details": str(e)}
                )
    
    except Exception as e:
        logger.error(
            "Failed to process presigned URL request",
            file_id=fileId,
            user_id=user_id,
            error=str(e),
            exc_info=True,
        )
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={"error": "Failed to process request", "details": str(e)}
        )
    
    finally:
        # Don't disconnect - pg_service is a singleton shared across requests
        pass


@router.get("/{fileId}/chunks", dependencies=[Depends(require_ingest_read)])
async def get_file_chunks(
    fileId: str,
    request: Request,
    limit: int = 100,
    offset: int = 0
):
    """
    Get text chunks for a file.
    
    Args:
        fileId: File identifier
        limit: Number of chunks to return (default: 100)
        offset: Offset for pagination (default: 0)
    
    Returns:
        List of text chunks with metadata
    """
    user_id = request.state.user_id
    
    from api.main import pg_service as postgres_service  # Use shared PostgresService instance
    # Connection is already established in startup
    
    try:
        async with postgres_service.acquire(request) as conn:
            # Check file access based on visibility and roles
            has_access, file_data, error_msg = await check_file_access(
                conn, uuid.UUID(fileId), user_id, request
            )
            
            if not has_access:
                status_code = status.HTTP_404_NOT_FOUND if error_msg == "File not found" else status.HTTP_403_FORBIDDEN
                return JSONResponse(
                    status_code=status_code,
                    content={"error": error_msg}
                )
            
            # Get chunk count
            chunk_count_row = await conn.fetchrow("""
                SELECT chunk_count FROM ingestion_files WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            # Get chunks
            chunks = await conn.fetch("""
                SELECT 
                    chunk_index,
                    text,
                    page_number,
                    char_offset,
                    token_count,
                    section_heading,
                    metadata,
                    created_at
                FROM ingestion_chunks
                WHERE file_id = $1
                ORDER BY chunk_index
                LIMIT $2 OFFSET $3
            """, uuid.UUID(fileId), limit, offset)
            
            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "fileId": fileId,
                    "total": chunk_count_row["chunk_count"] if chunk_count_row else 0,
                    "limit": limit,
                    "offset": offset,
                    "chunks": [
                        {
                            "chunkIndex": chunk["chunk_index"],
                            "text": chunk["text"],
                            "pageNumber": chunk["page_number"],
                            "charOffset": chunk["char_offset"],
                            "tokenCount": chunk["token_count"],
                            "sectionHeading": chunk["section_heading"],
                            "metadata": chunk["metadata"],
                            "createdAt": chunk["created_at"].isoformat(),
                        }
                        for chunk in chunks
                    ]
                }
            )
    
    except Exception as e:
        logger.error(
            "Failed to get file chunks",
            file_id=fileId,
            user_id=user_id,
            error=str(e),
            exc_info=True,
        )
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={"error": "Failed to retrieve chunks", "details": str(e)}
        )
    
    finally:
        # Don't disconnect - pg_service is a singleton shared across requests
        pass


@router.get("/{fileId}/vectors", dependencies=[Depends(require_ingest_read)])
async def get_file_vectors(
    fileId: str,
    request: Request,
    limit: int = 100,
    offset: int = 0
):
    """
    Get vector embeddings for a file's chunks from Milvus.
    
    Args:
        fileId: File identifier
        limit: Number of vectors to return (default: 100)
        offset: Offset for pagination (default: 0)
    
    Returns:
        List of chunk embeddings with metadata from Milvus
    """
    user_id = request.state.user_id
    
    from api.main import pg_service as postgres_service  # Use shared PostgresService instance
    # Connection is already established in startup
    
    try:
        async with postgres_service.acquire(request) as conn:
            # Check file access based on visibility and roles
            has_access, file_data, error_msg = await check_file_access(
                conn, uuid.UUID(fileId), user_id, request
            )
            
            if not has_access:
                status_code = status.HTTP_404_NOT_FOUND if error_msg == "File not found" else status.HTTP_403_FORBIDDEN
                return JSONResponse(
                    status_code=status_code,
                    content={"error": error_msg}
                )
            
            # Get vector count from PostgreSQL
            vector_count_row = await conn.fetchrow("""
                SELECT vector_count FROM ingestion_files WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            vector_count = vector_count_row["vector_count"] if vector_count_row else 0
            
            if vector_count == 0:
                return JSONResponse(
                    status_code=status.HTTP_404_NOT_FOUND,
                    content={"error": "No vectors available for this file"}
                )
            
            # Query Milvus for vectors
            # Note: Vectors are stored in Milvus, not PostgreSQL
            # We'll return metadata about vectors from chunks table
            chunks_with_vectors = await conn.fetch("""
                SELECT 
                    chunk_index,
                    text,
                    page_number,
                    token_count,
                    created_at
                FROM ingestion_chunks
                WHERE file_id = $1
                ORDER BY chunk_index
                LIMIT $2 OFFSET $3
            """, uuid.UUID(fileId), limit, offset)
            
            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "fileId": fileId,
                    "total": vector_count,
                    "limit": limit,
                    "offset": offset,
                    "note": "Vectors stored in Milvus - showing chunk metadata",
                    "chunks": [
                        {
                            "chunkIndex": chunk["chunk_index"],
                            "text": chunk["text"][:200] + "..." if len(chunk["text"]) > 200 else chunk["text"],
                            "pageNumber": chunk["page_number"],
                            "tokenCount": chunk["token_count"],
                            "hasVector": True,  # If in this result, it has a vector
                            "createdAt": chunk["created_at"].isoformat(),
                        }
                        for chunk in chunks_with_vectors
                    ]
                }
            )
    
    except Exception as e:
        logger.error(
            "Failed to get file vectors",
            file_id=fileId,
            user_id=user_id,
            error=str(e),
            exc_info=True,
        )
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={"error": "Failed to retrieve vectors", "details": str(e)}
        )
    
    finally:
        # Don't disconnect - pg_service is a singleton shared across requests
        pass


@router.get("/{fileId}/markdown", dependencies=[Depends(require_ingest_read)])
async def get_file_markdown(
    fileId: str,
    request: Request
):
    """
    Get markdown representation of a file from MinIO.
    
    Args:
        fileId: File identifier
    
    Returns:
        Markdown content with metadata
    """
    # Validate UUID format before any database operations
    try:
        file_uuid = uuid.UUID(fileId)
    except ValueError:
        return JSONResponse(
            status_code=status.HTTP_400_BAD_REQUEST,
            content={"error": "Invalid file ID format", "details": "File ID must be a valid UUID"}
        )
    
    user_id = request.state.user_id
    
    config = Config().to_dict()
    postgres_service = PostgresService(config, request)
    minio_service = MinIOService(config)
    await postgres_service.connect()
    
    try:
        async with postgres_service.acquire(request) as conn:
            # Check file access based on visibility and roles
            has_access, file_data, error_msg = await check_file_access(
                conn, file_uuid, user_id, request
            )
            
            if not has_access:
                status_code = status.HTTP_404_NOT_FOUND if error_msg == "File not found" else status.HTTP_403_FORBIDDEN
                return JSONResponse(
                    status_code=status_code,
                    content={"error": error_msg}
                )
            
            # Get markdown metadata from PostgreSQL
            markdown_row = await conn.fetchrow("""
                SELECT 
                    original_filename,
                    markdown_path,
                    has_markdown,
                    image_count,
                    created_at
                FROM ingestion_files
                WHERE file_id = $1
            """, file_uuid)
            
            if not markdown_row or not markdown_row["has_markdown"] or not markdown_row["markdown_path"]:
                return JSONResponse(
                    status_code=status.HTTP_404_NOT_FOUND,
                    content={"error": "Markdown not available for this file"}
                )
            
            # Fetch markdown content from MinIO
            import asyncio
            loop = asyncio.get_event_loop()
            markdown_content = await loop.run_in_executor(
                None,
                minio_service.get_file_content,
                markdown_row["markdown_path"]
            )
            
            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "fileId": fileId,
                    "filename": markdown_row["original_filename"],
                    "markdown": markdown_content,
                    "hasImages": markdown_row["image_count"] > 0 if markdown_row["image_count"] else False,
                    "imageCount": markdown_row["image_count"] or 0,
                    "length": len(markdown_content),
                    "createdAt": markdown_row["created_at"].isoformat(),
                }
            )
    
    except Exception as e:
        logger.error(
            "Failed to get file markdown",
            file_id=fileId,
            user_id=user_id,
            error=str(e),
            exc_info=True,
        )
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={"error": "Failed to retrieve markdown", "details": str(e)}
        )
    
    finally:
        # Don't disconnect - pg_service is a singleton shared across requests
        pass


class DocumentSearchRequest(BaseModel):
    """Request model for single-document search."""
    query: str = Field(..., description="Search query")
    limit: int = Field(default=10, ge=1, le=100, description="Max results to return")


@router.post("/{fileId}/search", dependencies=[Depends(require_ingest_read)])
async def search_within_document(
    fileId: str,
    request: Request,
    search_request: DocumentSearchRequest
):
    """
    Search within a single document using semantic search.
    
    Args:
        fileId: File identifier
        search_request: Search parameters
    
    Returns:
        List of matching chunks with relevance scores
    """
    user_id = request.state.user_id
    
    config = Config().to_dict()
    postgres_service = PostgresService(config, request)
    milvus_service = MilvusService(config)
    
    await postgres_service.connect()
    
    try:
        async with postgres_service.acquire(request) as conn:
            # Verify file exists and user owns it
            file_row = await conn.fetchrow("""
                SELECT user_id
                FROM ingestion_files
                WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            if not file_row:
                return JSONResponse(
                    status_code=status.HTTP_404_NOT_FOUND,
                    content={"error": "File not found"}
                )
            
            if str(file_row["user_id"]) != user_id:
                return JSONResponse(
                    status_code=status.HTTP_403_FORBIDDEN,
                    content={"error": "Unauthorized access"}
                )
            
            # Generate embedding for query
            from processors.embedder import Embedder
            embedder = Embedder(config)
            query_embedding = await embedder.embed_single(search_request.query)
            
            if not query_embedding:
                return JSONResponse(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    content={"error": "Failed to generate query embedding"}
                )
            
            # Search Milvus with file_id filter
            # Note: hybrid_search doesn't support custom filter_expr yet
            # So we'll get all results for the user and filter client-side
            all_results = milvus_service.hybrid_search(
                query_embedding=query_embedding,
                user_id=user_id,
                top_k=search_request.limit * 10,  # Get more results to filter
            )
            
            # Filter results to only this document
            search_results = [
                r for r in all_results 
                if r.get("file_id") == fileId
            ][:search_request.limit]
            
            # Get chunk details for results
            if search_results:
                # Build results directly from search results (they already have all needed fields)
                chunks = await conn.fetch("""
                    SELECT 
                        f.filename
                    FROM ingestion_files f
                    WHERE f.file_id = $1::uuid
                """, uuid.UUID(fileId))
                
                filename = chunks[0]["filename"] if chunks else "unknown"
                
                # Combine results
                results = []
                for result in search_results:
                    results.append({
                        "fileId": result["file_id"],
                        "filename": filename,
                        "chunkIndex": result["chunk_index"],
                        "pageNumber": result.get("page_number"),
                        "text": result["text"],
                        "score": result["score"],
                    })
                
                return JSONResponse(
                    status_code=status.HTTP_200_OK,
                    content={
                        "query": search_request.query,
                        "fileId": fileId,
                        "results": results,
                        "total": len(results),
                        "limit": search_request.limit,
                    }
                )
            else:
                return JSONResponse(
                    status_code=status.HTTP_200_OK,
                    content={
                        "query": search_request.query,
                        "fileId": fileId,
                        "results": [],
                        "total": 0,
                        "limit": search_request.limit,
                    }
                )
    
    except Exception as e:
        logger.error(
            "Failed to search within document",
            file_id=fileId,
            user_id=user_id,
            query=search_request.query,
            error=str(e),
            exc_info=True,
        )
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={"error": "Search failed", "details": str(e)}
        )
    
    finally:
        # Don't disconnect - pg_service is a singleton shared across requests
        pass


@router.delete("/{fileId}", dependencies=[Depends(require_ingest_delete)])
async def delete_file(fileId: str, request: Request):
    """
    Delete file and all associated data.
    
    Deletes:
        - File from MinIO storage
        - Vectors from Milvus
        - Encryption keys from AuthZ keystore
        - Metadata from PostgreSQL (cascades to chunks and status)
    """
    user_id = request.state.user_id
    
    config = Config().to_dict()
    postgres_service = PostgresService(config, request)
    minio_service = MinIOService(config)
    milvus_service = MilvusService(config)
    encryption_client = EncryptionClient(config)
    
    await postgres_service.connect()
    
    try:
        async with postgres_service.acquire(request) as conn:
            # Get file record
            file_row = await conn.fetchrow("""
                SELECT user_id, storage_path
                FROM ingestion_files
                WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            if not file_row:
                return JSONResponse(
                    status_code=status.HTTP_404_NOT_FOUND,
                    content={"error": "File not found"}
                )
            
            # Verify ownership
            if str(file_row["user_id"]) != user_id:
                return JSONResponse(
                    status_code=status.HTTP_403_FORBIDDEN,
                    content={"error": "Unauthorized access"}
                )
            
            storage_path = file_row["storage_path"]
            
            # Delete vectors from Milvus
            try:
                milvus_service.delete_file_vectors(fileId)
                logger.info("Deleted Milvus vectors", file_id=fileId)
            except Exception as e:
                logger.warning(
                    "Failed to delete vectors from Milvus",
                    file_id=fileId,
                    error=str(e),
                )
            
            # Delete file from MinIO
            try:
                await minio_service.delete_file(storage_path)
                logger.info("Deleted MinIO file", file_id=fileId, storage_path=storage_path)
            except Exception as e:
                logger.warning(
                    "Failed to delete file from MinIO",
                    file_id=fileId,
                    storage_path=storage_path,
                    error=str(e),
                )
            
            # Delete encryption keys from AuthZ keystore
            try:
                await encryption_client.delete_file_keys(fileId)
                logger.info("Deleted encryption keys", file_id=fileId)
            except Exception as e:
                logger.warning(
                    "Failed to delete encryption keys",
                    file_id=fileId,
                    error=str(e),
                )
            
            # Delete from PostgreSQL (cascades to status and chunks)
            await conn.execute("""
                DELETE FROM ingestion_files WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            logger.info(
                "File deleted successfully",
                file_id=fileId,
                user_id=user_id,
                storage_path=storage_path,
            )
            
            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "message": "File deleted successfully",
                    "fileId": fileId,
                }
            )
    
    except Exception as e:
        logger.error(
            "Failed to delete file",
            file_id=fileId,
            user_id=user_id,
            error=str(e),
            exc_info=True,
        )
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={"error": "Failed to delete file", "details": str(e)}
        )
    
    finally:
        # Don't disconnect - pg_service is a singleton shared across requests
        pass


@router.post("/{fileId}/move", dependencies=[Depends(require_ingest_write)])
async def move_file(fileId: str, request: Request):
    """
    Move a document between visibility modes and update role bindings.

    Body:
      - visibility: 'personal' or 'shared'
      - roleIds: [] required if visibility='shared'

    Rules:
      - Moving to personal is allowed only if actor is the owner.
      - Moving to shared requires roleIds.
    """
    body = await request.json()
    target_visibility = body.get("visibility")
    role_ids = body.get("roleIds", [])
    actor_id = request.state.user_id

    if target_visibility not in ("personal", "shared"):
        return JSONResponse(
            status_code=status.HTTP_400_BAD_REQUEST,
            content={"error": "visibility must be 'personal' or 'shared'"},
        )

    if target_visibility == "shared" and (not isinstance(role_ids, list) or len(role_ids) == 0):
        return JSONResponse(
            status_code=status.HTTP_400_BAD_REQUEST,
            content={"error": "roleIds required for shared visibility"},
        )

    from api.main import pg_service as postgres_service  # Use shared PostgresService instance
    # Connection is already established in startup

    # Fetch current record under RLS
    async with postgres_service.acquire(request) as conn:
        row = await conn.fetchrow(
            """
            SELECT file_id, owner_id, visibility
            FROM ingestion_files
            WHERE file_id = $1
            """,
            uuid.UUID(fileId),
        )

        if not row:
            return JSONResponse(status_code=status.HTTP_404_NOT_FOUND, content={"error": "File not found"})

        owner_id = str(row["owner_id"])
        current_visibility = row["visibility"]

        # Guard: shared -> personal requires ownership
        if target_visibility == "personal" and actor_id != owner_id:
            return JSONResponse(
                status_code=status.HTTP_403_FORBIDDEN,
                content={"error": "Only the owner can move a shared file to personal"},
            )

    # Update visibility + roles atomically
    try:
        await postgres_service.update_document_visibility_and_roles(
            file_id=fileId,
            visibility=target_visibility,
            role_ids=role_ids,
            actor_id=actor_id,
        )
        await postgres_service.insert_audit(
            actor_id=actor_id,
            action="file.move",
            resource_type="file",
            resource_id=fileId,
            details={
                "from_visibility": current_visibility,
                "to_visibility": target_visibility,
                "role_ids": role_ids,
            },
        )
    except Exception as exc:
        logger.error("Failed to move file", error=str(exc))
        return JSONResponse(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, content={"error": "Failed to move file"})

    return {"status": "ok", "visibility": target_visibility, "roleIds": role_ids}

@router.post("/{fileId}/reprocess", dependencies=[Depends(require_ingest_write)])
async def reprocess_file(fileId: str, request: Request):
    """
    Reprocess a document - optionally from a specific stage.
    
    This is useful when:
    - Chunking strategy has been updated
    - Embedding model has changed
    - Document processing failed partially
    - You want to regenerate embeddings
    
    Process:
    1. Verify file exists and user owns it
    2. Conditionally delete chunks/vectors based on start_stage
    3. Reset ingestion status to 'queued'
    4. Add job back to Redis queue for reprocessing
    
    Request Body (optional JSON):
        processing_config: Processing configuration including:
            - start_stage: Stage to start from (parsing, chunking, cleanup, markdown, embedding, indexing)
            - llm_cleanup_enabled, marker_enabled, etc.
    
    Returns:
        Success message with file_id
    """
    user_id = request.state.user_id
    
    # Parse optional processing config from request body
    processing_config = {}
    try:
        body = await request.json()
        processing_config = body.get("processing_config", {})
        if processing_config:
            logger.info(
                "Reprocess with custom processing config",
                file_id=fileId,
                config=processing_config,
            )
    except Exception:
        # No body or invalid JSON - use defaults
        pass
    
    # Determine start stage - affects what data to delete
    start_stage = processing_config.get("start_stage", "parsing")
    # Stages that require deleting chunks
    stages_needing_chunk_delete = ["parsing", "chunking"]
    # Stages that require deleting vectors
    stages_needing_vector_delete = ["parsing", "chunking", "cleanup", "embedding", "indexing"]
    
    config = Config().to_dict()
    postgres_service = PostgresService(config, request)
    milvus_service = MilvusService(config)
    
    await postgres_service.connect()
    
    try:
        async with postgres_service.acquire(request) as conn:
            # Get file record
            file_row = await conn.fetchrow("""
                SELECT user_id, filename, storage_path, mime_type, original_filename
                FROM ingestion_files
                WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            if not file_row:
                return JSONResponse(
                    status_code=status.HTTP_404_NOT_FOUND,
                    content={"error": "File not found"}
                )
            
            # Verify ownership
            if str(file_row["user_id"]) != user_id:
                return JSONResponse(
                    status_code=status.HTTP_403_FORBIDDEN,
                    content={"error": "Unauthorized access"}
                )
            
            filename = file_row["filename"]
            storage_path = file_row["storage_path"]
            mime_type = file_row["mime_type"]
            original_filename = file_row["original_filename"]
            
            logger.info(
                "Starting document reprocessing",
                file_id=fileId,
                filename=filename,
                mime_type=mime_type,
                user_id=user_id,
                start_stage=start_stage,
            )
            
            # Only delete chunks if starting from early stages
            if start_stage in stages_needing_chunk_delete:
                deleted_chunks = await conn.execute("""
                    DELETE FROM ingestion_chunks WHERE file_id = $1
                """, uuid.UUID(fileId))
                
                logger.info(
                    "Deleted existing chunks",
                    file_id=fileId,
                    chunks_deleted=deleted_chunks,
                )
            
            # Always clear processing history for stages we're re-running
            await conn.execute("""
                DELETE FROM processing_history WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            # Delete existing processing strategy results
            await conn.execute("""
                DELETE FROM processing_strategy_results WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            logger.info(
                "Cleared processing history",
                file_id=fileId,
            )
            
            # Only delete vectors if needed for this stage
            if start_stage in stages_needing_vector_delete:
                try:
                    milvus_service.delete_file_vectors(fileId)
                    logger.info("Deleted Milvus vectors", file_id=fileId)
                except Exception as e:
                    logger.warning(
                        "Failed to delete vectors from Milvus (may not exist)",
                        file_id=fileId,
                        error=str(e),
                    )
            
            # Reset ingestion status to 'queued'
            await conn.execute("""
                UPDATE ingestion_status
                SET 
                    stage = 'queued',
                    progress = 0,
                    chunks_processed = 0,
                    total_chunks = 0,
                    pages_processed = 0,
                    total_pages = 0,
                    error_message = NULL,
                    updated_at = NOW()
                WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            logger.info("Reset ingestion status", file_id=fileId)
            
            # Add job back to Redis queue
            redis_client = redis_async.Redis(
                host=config.get("redis_host", "localhost"),
                port=config.get("redis_port", 6379),
                decode_responses=True,
            )
            
            try:
                job_data = {
                    "job_id": fileId,  # Worker expects job_id field
                    "file_id": fileId,
                    "user_id": user_id,
                    "storage_path": storage_path,
                    "original_filename": original_filename or filename,
                    "mime_type": mime_type,  # Required for text extraction
                    "reprocess": "true",  # Flag to indicate this is a reprocess
                    "start_stage": start_stage,  # Stage to start processing from
                }
                
                # Include processing config if provided
                if processing_config:
                    job_data["processing_config"] = json.dumps(processing_config)
                
                stream_name = config.get("redis_stream", "jobs:ingestion")
                await redis_client.xadd(stream_name, job_data)
                
                logger.info(
                    "Added reprocess job to queue",
                    file_id=fileId,
                    stream=stream_name,
                )
            finally:
                await redis_client.aclose()
            
            logger.info(
                "Document reprocessing initiated",
                file_id=fileId,
                filename=filename,
                user_id=user_id,
            )
            
            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "message": f"Document reprocessing initiated from {start_stage}",
                    "fileId": fileId,
                    "filename": filename,
                    "status": "queued",
                    "start_stage": start_stage,
                }
            )
    
    except Exception as e:
        logger.error(
            "Failed to reprocess file",
            file_id=fileId,
            user_id=user_id,
            error=str(e),
            exc_info=True,
        )
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={"error": "Failed to reprocess file", "details": str(e)}
        )
    
    finally:
        # Don't disconnect - pg_service is a singleton shared across requests
        pass


@router.get("/{fileId}/export", dependencies=[Depends(require_ingest_read)])
async def export_file(
    fileId: str,
    request: Request,
    format: str = Query("markdown", regex="^(markdown|html|text|docx|pdf)$")
):
    """
    Export document by reconstructing from markdown chunks.
    
    Supports formats:
    - markdown: Original markdown with preserved structure
    - html: HTML conversion from markdown
    - text: Plain text (markdown stripped)
    - docx: Microsoft Word document
    - pdf: PDF document
    
    Args:
        fileId: File UUID
        format: Export format (markdown, html, text, docx, pdf)
    
    Returns:
        File content in requested format
    """
    user_id = request.state.user_id
    
    config = Config().to_dict()
    postgres_service = PostgresService(config, request)
    
    await postgres_service.connect()
    
    try:
        async with postgres_service.acquire(request) as conn:
            # Get file metadata and verify ownership
            file_row = await conn.fetchrow("""
                SELECT file_id, filename, original_filename, user_id
                FROM ingestion_files
                WHERE file_id = $1
            """, uuid.UUID(fileId))
            
            if not file_row:
                logger.warning("File not found for export", file_id=fileId, user_id=user_id)
                return JSONResponse(
                    status_code=status.HTTP_404_NOT_FOUND,
                    content={"error": "File not found"}
                )
            
            if str(file_row["user_id"]) != user_id:
                logger.warning("Unauthorized export attempt", file_id=fileId, user_id=user_id, owner_id=str(file_row["user_id"]))
                return JSONResponse(
                    status_code=status.HTTP_403_FORBIDDEN,
                    content={"error": "Unauthorized"}
                )
            
            filename = file_row["original_filename"] or file_row["filename"]
            
            # Get all chunks ordered by chunk_index
            chunk_rows = await conn.fetch("""
                SELECT chunk_index, text, page_number
                FROM ingestion_chunks
                WHERE file_id = $1
                ORDER BY chunk_index ASC
            """, uuid.UUID(fileId))
            
            if not chunk_rows:
                logger.warning("No chunks found for export", file_id=fileId, user_id=user_id)
                return JSONResponse(
                    status_code=status.HTTP_404_NOT_FOUND,
                    content={"error": "No content available for export"}
                )
            
            # Reconstruct document from chunks
            markdown_content = "\n\n".join([row["text"] for row in chunk_rows])
            
            logger.info(
                "Exporting document",
                file_id=fileId,
                filename=filename,
                format=format,
                chunk_count=len(chunk_rows),
                user_id=user_id,
            )
            
            # Convert to requested format
            if format == "markdown":
                content = markdown_content.encode("utf-8")
                media_type = "text/markdown"
                extension = "md"
            
            elif format == "html":
                # Convert markdown to HTML
                try:
                    import markdown
                    html_body = markdown.markdown(
                        markdown_content,
                        extensions=['extra', 'codehilite', 'tables', 'toc']
                    )
                    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{filename}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1, h2, h3 {{ margin-top: 24px; margin-bottom: 16px; }}
        code {{ background-color: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
        pre {{ background-color: #f4f4f4; padding: 16px; border-radius: 6px; overflow-x: auto; }}
        blockquote {{ border-left: 4px solid #ddd; padding-left: 16px; color: #666; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f4f4f4; }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""
                    content = html_content.encode("utf-8")
                    media_type = "text/html"
                    extension = "html"
                except ImportError:
                    logger.error("markdown library not available for HTML export")
                    return JSONResponse(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        content={"error": "HTML export not available"}
                    )
            
            elif format == "text":
                # Strip markdown formatting
                import re
                text_content = markdown_content
                # Remove markdown headers
                text_content = re.sub(r'^#+\s+', '', text_content, flags=re.MULTILINE)
                # Remove bold/italic
                text_content = re.sub(r'\*\*(.+?)\*\*', r'\1', text_content)
                text_content = re.sub(r'\*(.+?)\*', r'\1', text_content)
                text_content = re.sub(r'__(.+?)__', r'\1', text_content)
                text_content = re.sub(r'_(.+?)_', r'\1', text_content)
                # Remove links but keep text
                text_content = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text_content)
                # Remove inline code
                text_content = re.sub(r'`(.+?)`', r'\1', text_content)
                
                content = text_content.encode("utf-8")
                media_type = "text/plain"
                extension = "txt"
            
            elif format == "docx":
                # Convert markdown to DOCX
                try:
                    from docx import Document
                    from docx.shared import Pt, Inches
                    import re
                    
                    doc = Document()
                    
                    # Parse markdown and add to document
                    lines = markdown_content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # Headings
                        if line.startswith('# '):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith('## '):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith('### '):
                            doc.add_heading(line[4:], level=3)
                        # Lists
                        elif line.startswith('- ') or line.startswith('* '):
                            doc.add_paragraph(line[2:], style='List Bullet')
                        elif re.match(r'^\d+\.\s', line):
                            doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
                        # Regular paragraph
                        else:
                            doc.add_paragraph(line)
                    
                    # Save to bytes
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    content = docx_buffer.getvalue()
                    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    extension = "docx"
                except ImportError:
                    logger.error("python-docx library not available for DOCX export")
                    return JSONResponse(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        content={"error": "DOCX export not available"}
                    )
            
            elif format == "pdf":
                # Convert markdown to PDF
                try:
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.units import inch
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
                    from reportlab.lib.enums import TA_LEFT
                    import re
                    
                    pdf_buffer = io.BytesIO()
                    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
                    
                    styles = getSampleStyleSheet()
                    story = []
                    
                    # Parse markdown and add to PDF
                    lines = markdown_content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            story.append(Spacer(1, 0.2*inch))
                            continue
                        
                        # Escape HTML entities
                        line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        
                        # Headings
                        if line.startswith('# '):
                            story.append(Paragraph(line[2:], styles['Heading1']))
                        elif line.startswith('## '):
                            story.append(Paragraph(line[3:], styles['Heading2']))
                        elif line.startswith('### '):
                            story.append(Paragraph(line[4:], styles['Heading3']))
                        # Lists
                        elif line.startswith('- ') or line.startswith('* '):
                            story.append(Paragraph(f"• {line[2:]}", styles['BodyText']))
                        elif re.match(r'^\d+\.\s', line):
                            story.append(Paragraph(line, styles['BodyText']))
                        # Regular paragraph
                        else:
                            story.append(Paragraph(line, styles['BodyText']))
                    
                    doc.build(story)
                    content = pdf_buffer.getvalue()
                    media_type = "application/pdf"
                    extension = "pdf"
                except ImportError as e:
                    logger.error("reportlab library not available for PDF export", error=str(e))
                    return JSONResponse(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        content={"error": "PDF export not available"}
                    )
            
            # Generate filename
            base_filename = filename.rsplit('.', 1)[0] if '.' in filename else filename
            export_filename = f"{base_filename}.{extension}"
            
            logger.info(
                "Document exported successfully",
                file_id=fileId,
                filename=export_filename,
                format=format,
                size_bytes=len(content),
                user_id=user_id,
            )
            
            return Response(
                content=content,
                media_type=media_type,
                headers={
                    "Content-Disposition": f'attachment; filename="{export_filename}"'
                }
            )
    
    except Exception as e:
        logger.error(
            "Failed to export file",
            file_id=fileId,
            format=format,
            user_id=user_id,
            error=str(e),
            exc_info=True,
        )
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content={"error": "Failed to export file", "details": str(e)}
        )
    
    finally:
        # Don't disconnect - pg_service is a singleton shared across requests
        pass

